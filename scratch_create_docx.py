import subprocess
import sys

try:
    import docx
except ImportError:
    print("Installing python-docx library...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Gmail 邮箱一键授权使用指南")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "微软雅黑"

    doc.add_paragraph("\n如果您需要为系统新增一个新的谷歌邮箱账号（Token）以接收验证码，请按照本指南步骤在客户端进行一键网页授权。本操作不需要您的电脑安装任何开发环境（如 Python 等）。\n")

    # 步骤一
    h1 = doc.add_paragraph()
    r1 = h1.add_run("第一步：准备工作（复制文件到软件目录）")
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.name = "微软雅黑"

    p1 = doc.add_paragraph(
        "请将收到的以下两个文件复制粘贴到您原有的软件运行目录下（即与主程序 app.exe 存放在同一个文件夹内）：\n"
        "  1. authorize.bat （一键授权脚本）\n"
        "  2. gmail_authorize.exe （免安装授权主程序）\n\n"
        "【正确的目录结构示例如下】：\n"
        "您的软件文件夹/\n"
        "  ├── app.exe (主程序)\n"
        "  ├── gmail_credentials/ (已有的邮箱证书文件夹)\n"
        "  ├── gmail_authorize.exe (本次复制的程序) <-- 新增放这里\n"
        "  └── authorize.bat (本次复制的脚本) <-- 新增放这里"
    )
    p1.paragraph_format.left_indent = Inches(0.2)

    # 步骤二
    h2 = doc.add_paragraph()
    r2 = h2.add_run("第二步：运行一键授权脚本")
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.name = "微软雅黑"

    p2 = doc.add_paragraph(
        "1. 在文件夹中找到复制过来的 authorize.bat，双击鼠标运行它。\n"
        "2. 弹出的黑色命令行窗口会提示：\n"
        "   请输入需要新增授权的邮箱别名 (例如 gmail4):\n"
        "3. 请输入您想给这个邮箱设置的别名（例如您打算叫它 gmail4，直接输入 gmail4 即可，不要包含特殊字符），然后按下键盘上的回车键 (Enter)。"
    )
    p2.paragraph_format.left_indent = Inches(0.2)

    # 步骤三
    h3 = doc.add_paragraph()
    r3 = h3.add_run("第三步：在浏览器中登录并完成网页授权")
    r3.font.size = Pt(14)
    r3.font.bold = True
    r3.font.name = "微软雅黑"

    p3 = doc.add_paragraph(
        "1. 回车后，系统会自动在您的默认浏览器中打开一个谷歌账号登录授权页面。\n"
        "2. 请登录或选择您要进行授权的那个 Gmail 邮箱账号。\n"
        "3. 如果页面提示 “未经身份验证的应用” (Google hasn't verified this app)：\n"
        "   • 请点击页面左下角的 “高级 (Advanced)” 链接。\n"
        "   • 展开后，点击下方出现的 “转至...（不安全）/ Go to ... (unsafe)” 链接。\n"
        "4. 在接下来的权限确认页面中，请勾选“查看您的电子邮件 (gmail.readonly)” 权限（此为只读权限，仅用于安全拉取邮件及验证码，绝不泄露个人隐私）。\n"
        "5. 点击 “继续 (Continue)” 或 “允许 (Allow)” 完成授权。\n"
        "6. 网页显示 “The authentication flow has completed. You may close this window.” 时，即可关闭浏览器窗口。"
    )
    p3.paragraph_format.left_indent = Inches(0.2)

    # 步骤四
    h4 = doc.add_paragraph()
    r4 = h4.add_run("第四步：确认授权成功")
    r4.font.size = Pt(14)
    r4.font.bold = True
    r4.font.name = "微软雅黑"

    p4 = doc.add_paragraph(
        "回到软件目录下，您会发现 gmail_credentials 文件夹下已经自动生成了您刚刚设置的别名文件夹（如 gmail4），且内部已生成了 token.json 文件。\n"
        "这代表该邮箱的 Token 已经获取成功！现在您就可以在主配置中使用此账号别名来拉取邮件了。"
    )
    p4.paragraph_format.left_indent = Inches(0.2)

    doc.save("Gmail邮箱一键授权使用指南.docx")
    print("Word document generated successfully: Gmail邮箱一键授权使用指南.docx")

if __name__ == "__main__":
    main()
